import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import re
import codecs

# Ensure UTF-8 encoding to prevent 'charmap' errors
sys.stdout = codecs.getwriter("utf-8")(sys.stdout.buffer, 'strict')

# File paths
input_txt_file = sys.argv[1]  # e.g., 'temp.txt'
output_docx = r"C:\\Users\\HP\\Downloads\\n8n\\Updated_Resume.docx"

# Read temp.txt content (force UTF-8 decoding)
with open(input_txt_file, "r", encoding="utf-8", errors="replace") as f:
    updated_text = f.read()

# Remove markdown-like formatting (** and ##) and the 'plaintext' block syntax
updated_text = re.sub(r"[*#]+", "", updated_text)
updated_text = re.sub(r"```plaintext", "", updated_text)  # Remove plaintext markers
updated_text = re.sub(r"'''$", "", updated_text, flags=re.MULTILINE)  # Remove ending three single quotes

# Fix spacing issues for words merged incorrectly
updated_text = re.sub(r"crossfunctional", "cross-functional", updated_text)
updated_text = re.sub(r"thirdparty", "third-party", updated_text)
updated_text = re.sub(r"highquality", "high-quality", updated_text)
updated_text = re.sub(r"decisionmaking", "decision-making", updated_text)
updated_text = re.sub(r"realtime", "real-time", updated_text)
updated_text = re.sub(r"endtoend", "end-to-end", updated_text)
updated_text = re.sub(r"datadriven", "data-driven", updated_text)

# Create a new Word document with adjusted margins
doc = Document()
sections = doc.sections
for section in sections:
    section.left_margin = Inches(1.0)  # Adjusted for better readability
    section.right_margin = Inches(1.0)

# Define styling functions
def add_heading(doc, text):
    """Adds a bold section heading with minimal spacing."""
    para = doc.add_paragraph()
    para.space_after = Pt(1)
    para.space_before = Pt(1)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return para

def add_subheading(doc, text):
    """Adds a bold subheading for job roles or institutions."""
    para = doc.add_paragraph()
    para.space_after = Pt(0)
    para.space_before = Pt(0)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return para

def add_date(doc, text):
    """Adds date text without bullet points."""
    para = doc.add_paragraph()
    para.space_after = Pt(0)
    para.space_before = Pt(0)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(50, 50, 50)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return para

def add_bullet_point(doc, text):
    """Adds bullet points for work experience details with indentation and removes duplicate bullets."""
    text = text.lstrip("• ")  # Remove extra bullets from text
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.3)  # Indent bullet points only
    para.space_after = Pt(0)
    para.space_before = Pt(0)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(50, 50, 50)  # Dark gray
    return para

# Define sections to structure resume
SECTIONS_TO_UPDATE = ["EXPERIENCE", "EDUCATION", "PROJECTS", "SKILLS", "LANGUAGES"]

def extract_sections(text):
    """Extracts sections from temp.txt and organizes them into a dictionary."""
    sections = {}
    current_section = None
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.upper() in SECTIONS_TO_UPDATE:
            current_section = line.upper()
            sections[current_section] = []
        elif current_section:
            sections[current_section].append(line)
    return sections

# Process the text and structure it
new_content = extract_sections(updated_text)

# Add name and contact details (First three lines of temp.txt)
lines = updated_text.split("\n")
if len(lines) > 0:
    doc.add_paragraph().add_run(lines[0].strip()).bold = True  # Name
if len(lines) > 1:
    doc.add_paragraph().add_run(lines[1].strip()).italic = True  # Job Title
if len(lines) > 2:
    doc.add_paragraph().add_run(lines[2].strip()).bold = False  # Contact Info

# Add a separator line
doc.add_paragraph("\n")

# Process sections
for section, content in new_content.items():
    add_heading(doc, section)
    for line in content:
        if "—" in line:  # If it's a job title or institution
            add_subheading(doc, line)
        elif re.match(r"^(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC) ", line):
            add_date(doc, line)  # Keep dates unbulleted
        elif "-" in line:  # Convert dashes to bullet points for descriptions
            add_bullet_point(doc, "• " + line.replace("-", "").strip())
        else:
            add_bullet_point(doc, "• " + line)  # Ensures bullet points for work descriptions

# Save formatted resume
doc.save(output_docx)

# Fix UnicodeEncodeError during printing
print("Formatted Resume saved at:", output_docx.encode("utf-8", "ignore").decode("utf-8"))
